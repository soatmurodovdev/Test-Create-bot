<<<<<<< HEAD
# Murodjon Soatmurodov tomonidan yaratilgan
=======
>>>>>>> 043a65b44659b22132508b88491bc6f3483f3db5
"""
Test natijalarini (savollar / javoblar) chiroyli, brendlangan ko'rinishda
chiqarish uchun modul. Uchta format qo'llab-quvvatlanadi: matn, DOCX, PDF.

Dizayn: binafsha sarlavha, rangli "pill" info-chiziq (daraja/soni/sana),
har savol chap tomonida rangli chiziq + och fon, pastda brend chizig'i.

MUHIM (moslik uchun): DOCX'da haqiqiy header/footer va floating (suv belgisi)
rasm ISHLATILMAYDI — chunki bu ba'zi mobil ilovalarda (Google Docs va h.k.)
noto'g'ri yoki umuman ko'rinmay chiqishi aniqlangan. Sarlavha va footer
oddiy hujjat tanasidagi jadval/paragraf sifatida qo'shiladi — bu barcha
ilovalarda bir xil ko'rinadi.
"""
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

# --- Brend ranglari ---
PRIMARY = "5B21B6"
ACCENT = "F59E0B"
GREEN = "10B981"
LIGHT_BG = "F3E8FF"
DARK_TEXT = "1E1B2E"
GRAY = "6B7280"
WHITE = "FFFFFF"

BOT_USERNAME = "@TestCreateAi_bot"
BRAND_LINE = f"🤖  Ushbu test {BOT_USERNAME}  sun'iy intellekti tomonidan yaratildi"

DIFFICULTY_LABELS = {
    "easy": {"uz": "OSON", "ru": "ЛЁГКИЙ", "en": "EASY"},
    "medium": {"uz": "O'RTA", "ru": "СРЕДНИЙ", "en": "MEDIUM"},
    "hard": {"uz": "QIYIN", "ru": "СЛОЖНЫЙ", "en": "HARD"},
}


# ---------------------------------------------------------------------------
# DOCX past-darajali (OXML) yordamchilari
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _paragraph_left_border(paragraph, hex_color, size=24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _paragraph_top_border(paragraph, hex_color, size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), hex_color)
    pBdr.append(top)
    pPr.append(pBdr)


def _table_fixed_layout(table, total_width_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = tbl.find(qn("w:tblGrid"))
    # ustun kengliklarini teng bo'lib qayta yozamiz
    cols = list(grid)
    n = len(cols)
    for gc in cols:
        grid.remove(gc)
    for _ in range(n):
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(total_width_dxa // n))
        grid.append(gridCol)


def _set_cell_width(cell, dxa):
    cell.width = Cm(dxa / 566.93)  # taxminiy dxa->cm (1 cm = 566.93 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(dxa))


PAGE_CONTENT_DXA = 10106  # A4, 1.5cm hoshiya


def _title_banner(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, PAGE_CONTENT_DXA)
    _shade_cell(cell, PRIMARY)
    cell.margin_top = Pt(10)
    cell.margin_bottom = Pt(10)
    cell.margin_left = Pt(12)

    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"🧠  {title}")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor.from_string(WHITE)

    if subtitle:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string("E9D5FF")

    _table_fixed_layout(table, PAGE_CONTENT_DXA)


def _info_strip(doc, pills):
    """pills: list of (text, hex_color) — 2 yoki 3 ta bo'lishi mumkin"""
    n = len(pills)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = PAGE_CONTENT_DXA // n
    for i, (text, color) in enumerate(pills):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, col_w)
        _shade_cell(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    _table_fixed_layout(table, PAGE_CONTENT_DXA)


def _footer_brand(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_top_border(p, ACCENT)
    r = p.add_run(BRAND_LINE)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def _meta_pills(meta, lang="uz"):
    diff_key = meta.get("difficulty", "medium")
    diff_label = DIFFICULTY_LABELS.get(diff_key, DIFFICULTY_LABELS["medium"]).get(lang, "O'RTA")
    pills = [
        (f"🎯 DARAJA: {diff_label}", ACCENT),
        (f"🔢 SAVOLLAR: {meta.get('count', len(meta.get('questions', [])))}", PRIMARY),
    ]
    if meta.get("date"):
        pills.append((f"📅 {meta['date']}", GREEN))
    return pills


# ---------------------------------------------------------------------------
# DOCX — faqat savollar
# ---------------------------------------------------------------------------

def build_questions_docx(questions, path, meta):
    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _title_banner(doc, "TEST TUZUVCHI AI", meta.get("subtitle", ""))
    doc.add_paragraph()
    _info_strip(doc, _meta_pills(meta, meta.get("lang", "uz")))

    letters = ["A", "B", "C", "D"]
    for i, q in enumerate(questions, start=1):
        qp = doc.add_paragraph()
        _shade_paragraph(qp, LIGHT_BG)
        _paragraph_left_border(qp, PRIMARY)
        r = qp.add_run(f"  {i}.  {q['question']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(DARK_TEXT)
        qp.paragraph_format.space_before = Pt(12)
        qp.paragraph_format.space_after = Pt(4)

        for k, v in q["options"].items():
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(0.8)
            op.paragraph_format.space_after = Pt(2)
            rk = op.add_run(f"{k}) ")
            rk.bold = True
            rk.font.color.rgb = RGBColor.from_string(PRIMARY)
            rv = op.add_run(v)
            rv.font.color.rgb = RGBColor.from_string(DARK_TEXT)

    doc.add_paragraph()
    _footer_brand(doc)
    doc.save(path)


# ---------------------------------------------------------------------------
# DOCX — faqat javoblar
# ---------------------------------------------------------------------------

def build_answers_docx(questions, path, meta):
    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _title_banner(doc, "JAVOBLAR", meta.get("subtitle", ""))
    doc.add_paragraph()
    _info_strip(doc, _meta_pills(meta, meta.get("lang", "uz")))
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parts_per_line = 5
    for i, q in enumerate(questions, start=1):
        r = p.add_run(f"{i}) {q['correct']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(PRIMARY)
        r2 = p.add_run("      " if i % parts_per_line else "\n")

    doc.add_paragraph()
    _footer_brand(doc)
    doc.save(path)


# ---------------------------------------------------------------------------
# PDF (reportlab) — savollar yoki javoblar
# ---------------------------------------------------------------------------

import re

_EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"
    "\U00002600-\U000027BF"
    "\U0001F1E6-\U0001F1FF"
    "\U00002190-\U000021FF"
    "\uFE0F"
    "]+",
    flags=re.UNICODE,
)


def _no_emoji(text):
    return _EMOJI_RE.sub("", text).strip()


def _hex(c):
    return colors.HexColor(f"#{c}")


def build_pdf(questions, path, meta, mode="questions"):
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        topMargin=1.3 * cm, bottomMargin=1.3 * cm,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "TitleBanner", parent=styles["Normal"], textColor=colors.white,
        fontSize=16, fontName="Helvetica-Bold", leading=20,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleBanner", parent=styles["Normal"], textColor=colors.HexColor("#E9D5FF"),
        fontSize=10, fontName="Helvetica-Oblique",
    )
    title_text = "TEST TUZUVCHI AI" if mode == "questions" else "JAVOBLAR"
    banner_cell = [Paragraph(title_text, title_style)]
    if meta.get("subtitle"):
        banner_cell.append(Paragraph(_no_emoji(meta["subtitle"]), subtitle_style))
    banner = Table([[banner_cell]], colWidths=[doc.width])
    banner.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _hex(PRIMARY)),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(banner)
    story.append(Spacer(1, 10))

    pills = _meta_pills(meta, meta.get("lang", "uz"))
    pill_style = ParagraphStyle(
        "Pill", parent=styles["Normal"], textColor=colors.white,
        fontSize=9, fontName="Helvetica-Bold", alignment=TA_CENTER,
    )
    pill_cells = [[Paragraph(_no_emoji(text), pill_style)] for text, _ in pills]
    pill_table = Table([[c[0] for c in pill_cells]], colWidths=[doc.width / len(pills)] * len(pills))
    pill_style_cmds = [
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    for i, (_, color) in enumerate(pills):
        pill_style_cmds.append(("BACKGROUND", (i, 0), (i, 0), _hex(color)))
    pill_table.setStyle(TableStyle(pill_style_cmds))
    story.append(pill_table)
    story.append(Spacer(1, 14))

    if mode == "questions":
        q_style = ParagraphStyle(
            "Question", parent=styles["Normal"], textColor=colors.HexColor(f"#{DARK_TEXT}"),
            fontSize=12, fontName="Helvetica-Bold", backColor=colors.HexColor(f"#{LIGHT_BG}"),
            leftIndent=6, borderPadding=(6, 6, 6, 6),
        )
        opt_style = ParagraphStyle(
            "Option", parent=styles["Normal"], textColor=colors.HexColor(f"#{DARK_TEXT}"),
            fontSize=11, leftIndent=18, spaceAfter=2,
        )
        for i, q in enumerate(questions, start=1):
            story.append(Spacer(1, 8))
            qt = Table([[Paragraph(f"<b>{i}.</b> {q['question']}", q_style)]], colWidths=[doc.width])
            qt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), _hex(LIGHT_BG)),
                ("LINEBEFORE", (0, 0), (0, 0), 3, _hex(PRIMARY)),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ]))
            story.append(qt)
            for k, v in q["options"].items():
                story.append(Paragraph(f"<b><font color='#{PRIMARY}'>{k})</font></b> {v}", opt_style))
    else:
        ans_style = ParagraphStyle(
            "Answer", parent=styles["Normal"], textColor=colors.HexColor(f"#{PRIMARY}"),
            fontSize=12, fontName="Helvetica-Bold",
        )
        rows, row = [], []
        for i, q in enumerate(questions, start=1):
            row.append(f"{i}. {q['correct']}")
            if len(row) == 5:
                rows.append(row)
                row = []
        if row:
            rows.append(row)
        for row_items in rows:
            story.append(Paragraph("&nbsp;&nbsp;&nbsp;&nbsp;".join(row_items), ans_style))
            story.append(Spacer(1, 4))

    story.append(Spacer(1, 20))
    brand_style = ParagraphStyle(
        "Brand", parent=styles["Normal"], textColor=colors.HexColor(f"#{GRAY}"),
        fontSize=9, alignment=TA_CENTER, borderWidth=0.5, borderColor=colors.HexColor(f"#{ACCENT}"),
        borderPadding=6,
    )
    story.append(Paragraph(_no_emoji(BRAND_LINE), brand_style))

    doc.build(story)


# ---------------------------------------------------------------------------
# Matn (Telegram xabar) — savollar yoki javoblar
# ---------------------------------------------------------------------------

def build_questions_text(questions, meta, tr_label="Savollar"):
    diff_key = meta.get("difficulty", "medium")
    diff_label = DIFFICULTY_LABELS.get(diff_key, DIFFICULTY_LABELS["medium"]).get(meta.get("lang", "uz"), "O'RTA")
    lines = [f"📝 {tr_label} — {diff_label}", ""]
    for i, q in enumerate(questions, start=1):
        lines.append(f"{i}. {q['question']}")
        for k, v in q["options"].items():
            lines.append(f"   {k}) {v}")
        lines.append("")
    lines.append(f"🤖 {BOT_USERNAME}")
    return "\n".join(lines)


def build_answers_text(questions, meta, tr_label="Javoblar"):
    lines = [f"✅ {tr_label}", ""]
    parts = [f"{i}️⃣{q['correct']}" for i, q in enumerate(questions, start=1)]
    for start in range(0, len(parts), 8):
        lines.append("   ".join(parts[start:start + 8]))
    lines.append("")
    lines.append(f"🤖 {BOT_USERNAME}")
    return "\n".join(lines)
