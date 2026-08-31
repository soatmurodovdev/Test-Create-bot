import docx


def build_quiz_docx(questions, path, title="Test"):
    d = docx.Document()
    d.add_heading(title, level=1)

    for i, q in enumerate(questions, start=1):
        d.add_paragraph(f"{i}. {q['question']}")
        for key, val in q["options"].items():
            d.add_paragraph(f"   {key}) {val}")
        d.add_paragraph("")

    d.add_page_break()
    d.add_heading("Javoblar", level=2)
    answers = ", ".join(f"{i + 1}-{q['correct']}" for i, q in enumerate(questions))
    d.add_paragraph(answers)

    d.save(path)
